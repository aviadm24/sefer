import PySimpleGUI as psg
from docx import Document
# set the theme for the screen/window
psg.theme('SandyBeach')
# define layout
layout = [[psg.Text('Choose Messsage1', size=(20, 1), font='Lucida', justification='left')],
          [psg.Combo(['#200', '#201', '#2200', '#2201'],
                     default_value='#200', key='board')],
          [psg.Text('Choose Messsage2 ', size=(30, 1), font='Lucida', justification='left')],
          [psg.Combo(['Vsm_Pyld_Fast_Cyclic', 'Vsm_Pyld_Slow_Cyclic', 'Vsm_Pyld_M0'],
                     key='dest')],
          [psg.Text('Choose additional Elements', size=(30, 1), font='Lucida', justification='left')],
          [psg.Listbox(values=['AddressesSensor', 'GimbalsOn', 'FlirOn', 'SwirOn'],
                       select_mode='extended', key='fac', size=(30, 6))],
          [psg.Listbox(values=['FOV_ON', 'LaserCode'],
                       select_mode='extended', key='fac2', size=(30, 6))],
          [psg.Button('SAVE', font=('Times New Roman', 12)), psg.Button('CANCEL', font=('Times New Roman', 12))]]
# Define Window
win = psg.Window('Customise your Use Case', layout)
# Read  values entered by user
e, v = win.read()
# close first window
win.close()
# access the selected value in the list box and add them to a string
strx = ""
for val in v['fac']:
    strx = strx + " " + val + ","

# display string in a popup
psg.popup('Options Chosen',
          'You will Travel from :' + v['board'] + ' to ' + v['dest'] + ' \nYour additional facilities are:' +
          strx[1:len(strx) - 1])
doc = Document()

# Add a Title to the document
doc.add_heading('MALAT USE CASE', 0)


# Table data in a form of list
data = (
    (v['board'], v['dest']),
    (v['fac'][0], v['fac2'][0]),
    (v['fac'][1], v['fac2'][1])
)

# Creating a table object
table = doc.add_table(rows=1, cols=2)

# Adding heading in the 1st row of the table
row = table.rows[0].cells
row[0].text = 'DLI'
row[1].text = 'UPL'

# Adding data from the list to the table
for id, name in data:
    # Adding a row and then adding data in it.
    row = table.add_row().cells
    row[0].text = str(id)
    row[1].text = name

# Adding style to a table
table.style = 'Colorful List'

# Now save the document to a location
doc.save('gfg.docx')

